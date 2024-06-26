import asyncio
import sys

from typing import Any
from dataclasses import dataclass
import docx
from PyQt5 import QtCore
from PyQt5.QtCore import QSize
from PyQt5.QtWidgets import (QApplication, QFileDialog, QLabel, QPlainTextEdit, QWidget, QMessageBox, QPushButton,
                             QComboBox, QScrollArea, QGridLayout)

import docx_ed.cfg as c
# from constants import READ_ONLY
from docx_ed.async_docx import FileManager
from docx_ed.file_reader import FileReader


# GOST_FILE = "GOSTs.json"


class SecondWindow(QWidget):
    selected_gost = 'new_gost'
    def __init__(self, main_window):
        super().__init__()

        #   Начальные значения
        self.pathFile = ''
        self.plain_text = None
        self.main_window = main_window
        self.setWindowTitle("Проверка файла")
        self.setGeometry(200, 40, 1440, 1024)
        self.setMaximumSize(QSize(1940, 990))
        self.setMinimumSize(QSize(1024, 700))

        self.initUI()

    def initUI(self):

        # Виджеты (элементы)
        self.title = QLabel("Проверка файла по ГОСТам", self)
        # self.title.setAlignment(QtCore.Qt.AlignCenter)
        self.title.setStyleSheet('''
                                QLabel {
                                          font-size: 29px;
                                          font-weight: 700;
                                          font-family: 'Aleo';
                                          color: #FFFFFF;
                                          width: 30%; height: 60%;
                                          padding: 0px 20px 0px 0px;
                                        }
                                    ''')

        # Кнопка открытия главного окна со своими настройками
        self.mainWindow_button = QPushButton("Проверить по своим параметрам", self)
        self.mainWindow_button.setMinimumSize(360, 80)
        self.mainWindow_button.setStyleSheet('''
                                QPushButton {
                                              font-weight: 700;
                                              background-color: #E9E9E9;
                                              font-family: 'Aleo';
                                              font-size: 20px;
                                              color: #000000;
                                              width: 50%; height: 70%;
                                              padding: 0px 10px 0px 10px
                                            }
                                    ''')

        self.labelGost = QLabel("Выберите ГОСТ", self)
        self.labelGost.setStyleSheet('''
                                    QLabel {
                                            font-size: 20px;
                                            font-family: 'Aleo';
                                            color: #F4F2F2;
                                            margin: 30px 30px 0px 0px;
                                            padding: 30px 0 10px 1px;
                                            }
                                        ''')

        self.gostPicked = QLabel(self)
        self.gostPicked.setVisible(True)
        self.gostPicked.setStyleSheet('''
                                    QLabel {
                                                font-size: 20px;
                                                font-family: 'Aleo';
                                                color: #F4F2F2;
                                            }
                                        ''')

        self.choiceGost = QComboBox(self)
        self.gost_keys = list(FileReader.get_actual_pre_gosts().keys())
        # self.gost_data = FileReader.read_file()    # Загрузка данных из JSON файла с использованием функции
        # self.gost_keys = list(self.gost_data.keys())   # Получение списка ключей
        self.choiceGost.addItems(self.gost_keys)  # Добавление ключей в QComboBox
        self.choiceGost.setStyleSheet('''
                     QComboBox {
                                    font-size: 20px;
                                    font-weight: 400;
                                    font-family: 'Aleo';
                                    color: #000000;
                                    background-color: #FFFFFF;
                                    margin: 0px 30px 0px 0px;
                                    border: 2px solid transparent; /* Прозрачная рамка */
                                    border-radius: 6px;
                                    padding: 0px 0px 0px 10px; /* Внутренний отступ */
                                    width: 30%; height: 60%;
                                }
           QComboBox::drop-down {
                                    subcontrol-position: right;
                                    width: 25%;
                                    /*padding: 1px;*/
                                }
    QComboBox QAbstractItemView {
                                    font-size: 20px;
                                    font-weight: 400;
                                    font-family: 'Aleo';
                                    color: #000000;
                                    background-color: #FFFFFF;
                                    margin: 0px 30px 0px 0px;
                                    selection-background-color: #C0C0C0; /* Цвет фона при наведении на элемент списка */
                                    selection-color: #000000; /* Цвет текста при наведении на элемент списка */
                                    /*padding: 30px 5px;*/
                                }
    QComboBox QAbstractItemView::item {
                                    padding: 30px 30px;
                                }
                            ''')

        self.elementPicked = QLabel(self)
        self.elementPicked.setVisible(True)
        self.elementPicked.setStyleSheet('''
                                            QLabel {
                                                        font-size: 20px;
                                                        font-family: 'Aleo';
                                                        color: #F4F2F2;
                                                    }
                                                ''')

        self.choiceElement = QComboBox(self)
        objElement = FileManager(1, docx.Document('../test2.docx'), 'tur', gost=self.selected_gost, doc_rej=False)
        asyncio.run(objElement.update_params_from_gost())
        styles = objElement.styles
        self.element_keys = list(styles.keys())
        self.choiceElement.addItems(self.element_keys)  # Добавление ключей в QComboBox
        self.choiceElement.setStyleSheet('''
                             QComboBox {
                                            font-size: 20px;
                                            font-weight: 400;
                                            font-family: 'Aleo';
                                            color: #000000;
                                            background-color: #FFFFFF;
                                            margin: 0px 30px 0px 0px;
                                            border: 2px solid transparent; /* Прозрачная рамка */
                                            border-radius: 6px;
                                            padding: 0px 0px 0px 10px; /* Внутренний отступ */
                                            width: 30%; height: 60%;
                                        }
                   QComboBox::drop-down {
                                            subcontrol-position: right;
                                            width: 25%;
                                            /*padding: 1px;*/
                                        }
            QComboBox QAbstractItemView {
                                            font-size: 20px;
                                            font-weight: 400;
                                            font-family: 'Aleo';
                                            color: #000000;
                                            background-color: #FFFFFF;
                                            margin: 0px 30px 0px 0px;
                                            selection-background-color: #C0C0C0; /* Цвет фона при наведении на элемент списка */
                                            selection-color: #000000; /* Цвет текста при наведении на элемент списка */
                                            /*padding: 30px 5px;*/
                                        }
            QComboBox QAbstractItemView::item {
                                            padding: 30px 30px;
                                        }
                                    ''')

        # Вывод параметров гостов
        self.fontStyleLabel = QLabel("Шрифт:", self)
        self.fontStyleLabel.setStyleSheet('''
                                          QLabel {
                                                  font-size: 20px;
                                                  font-weight: 400;
                                                  font-family: 'Aleo';
                                                  color: #F4F2F2;
                                                  /* border: 1px solid red; */
                                                  margin: 0px 30px 0px 30px;
                                                  padding: 36px 0px 36px 0px;
                                                }
                                            ''')

        self.fontSizeLabel = QLabel("Размер шрифта:", self)
        self.fontSizeLabel.setStyleSheet('''
                                          QLabel {
                                                  font-size: 20px;
                                                  font-weight: 400;
                                                  font-family: 'Aleo';
                                                  color: #F4F2F2;
                                                  /* border: 1px solid red; */
                                                  margin: 0px 30px 0px 30px;
                                                  
                                                }
                                            ''')

        self.paragraphIndentLabel = QLabel("Абзацный отступ:", self)
        self.paragraphIndentLabel.setStyleSheet('''
                                          QLabel {
                                                  font-size: 20px;
                                                  font-weight: 400;
                                                  font-family: 'Aleo';
                                                  color: #F4F2F2;
                                                  /* border: 1px solid red; */
                                                  margin: 0px 30px 0px 30px;
                                                }
                                            ''')

        self.intervalLabel = QLabel("Межстрочный интервал:", self)
        self.intervalLabel.setStyleSheet('''
                                          QLabel {
                                                  font-size: 20px;
                                                  font-weight: 400;
                                                  font-family: 'Aleo';
                                                  color: #F4F2F2;
                                                  /* border: 1px solid red; */
                                                  margin: 0px 30px 0px 30px;
                                                }
                                            ''')

        self.alignmentLabel = QLabel("Выравнивание:", self)
        self.alignmentLabel.setStyleSheet('''
                                          QLabel {
                                                  font-size: 20px;
                                                  font-weight: 400;
                                                  font-family: 'Aleo';
                                                  color: #F4F2F2;
                                                  /* border: 1px solid red; */
                                                  margin: 0px 30px 0px 30px;
                                                }
                                            ''')

        self.styleLabel = QLabel("Название стиля:", self)
        self.styleLabel.setStyleSheet('''
                                                  QLabel {
                                                          font-size: 20px;
                                                          font-weight: 400;
                                                          font-family: 'Aleo';
                                                          color: #F4F2F2;
                                                          /* border: 1px solid red; */
                                                          margin: 0px 30px 0px 30px;
                                                        }
                                                    ''')

        # Кнопка выбора файла
        self.pickFileButton = QPushButton("Выбрать файл (docx)", self)
        self.pickFileButton.setMinimumSize(600, 70)
        self.pickFileButton.setAcceptDrops(True)
        self.pickFileButton.setStyleSheet('''
                            QPushButton {
                                          font-weight: 400;
                                          font-family: 'Aleo';
                                          font-size: 20px;
                                          color: #FFFFFF;
                                          text-align: center;
                                          border: 3px solid #FFFFFF;
                                          border-radius: 36px;
                                          padding: 10px 30px;
                                          margin: 55px 170px 0px 170px;
                                          width: 30%; height: 50%;
                                    }
                                ''')

        self.filePicked = QLabel("", self)
        self.filePicked.setStyleSheet('''
                                          QLabel {
                                                  font-size: 20px;
                                                  font-weight: 400;
                                                  font-family: 'Aleo';
                                                  color: #F4F2F2;
                                                }
                                            ''')
        self.filePicked.setAlignment(QtCore.Qt.AlignCenter)

        self.checkFile = QPushButton("Проверить файл", self)
        self.checkFile.setMinimumSize(600, 80)
        self.checkFile.setStyleSheet('''
                        QPushButton {
                                      font-weight: 400;
                                      font-family: 'Aleo';
                                      font-size: 20px;
                                      color: #FFFFFF;
                                      text-align: center;
                                      border: 3px solid #FFFFFF;
                                      border-radius: 36px;
                                      padding: 10px 30px;
                                      margin: 30px 170px 35px 170px;
                                      width: 30%; height: 50%;
                                    }
                                ''')

        self.answer = QScrollArea(self)
        self.answer.setMinimumSize(900, 80)
        self.answer.setWidgetResizable(True)
        # self.answer.setAlignment(QtCore.Qt.AlignLeading | QtCore.Qt.AlignLeft | QtCore.Qt.AlignTop)
        self.answer.setStyleSheet('''
                                            QScrollArea {
                                                    background-color: #FFFFFF;
                                                    margin-right: 0px;
                                                    width: 500%;
                                            }
                                        ''')

        self.plain_text = QPlainTextEdit()
        self.plain_text.setReadOnly(True)
        self.plain_text.setMinimumSize(300, 80)
        self.plain_text.setStyleSheet('''
                                    QPlainTextEdit {
                                            background-color: #FFFFFF;
                                            border: 4px solid black;
                                            border-radius: 50px;
                                            padding: 25px 10px 25px 40px;
                                            color: #000000;
                                            font-size: 20px;
                                            font-weight: 400;
                                            margin: 20px 40px;
                                            width: 500%;
                                    }
                                ''')

        self.downloadFile = QPushButton("Скачать проверенный\nфайл", self)
        self.downloadFile.setMinimumSize(350, 80)
        self.downloadFile.setStyleSheet('''
                            QPushButton {
                                            font-weight: 700;
                                            background-color: #E9E9E9;
                                            font-family: 'Aleo';
                                            font-size: 20px;
                                            color: #000000;
                                            border: 3px solid #FFFFFF;
                                            border-radius: 36px;
                                            margin: 30px 50px 35px 50px;
                                            padding: 10px 0px;
                                            width: 70%; height: 50%;
                                        }
                                    ''')

        # layout = QVBoxLayout(self)
        # layout.addWidget(self.plain_text)
        #
        # w = QWidget()
        # w.setLayout(layout)
        # self.answer.setWidget(w)
        # w.setStyleSheet("background-color: #0074BA;")

        #   Обработка событий
        self.choiceGost.activated.connect(self.get_params_from_ghost)
        self.choiceElement.activated.connect(self.get_params_from_element)
        self.gostPicked.setText(self.choiceGost.currentText())
        self.elementPicked.setText(self.choiceElement.currentText())
        self.pickFileButton.clicked.connect(self.pickFileButton_Clicked)
        self.checkFile.clicked.connect(self.checkFile_Clicked)
        self.downloadFile.clicked.connect(self.save_ready_file)
        self.mainWindow_button.clicked.connect(self.open_main_window)  # Открытие главного окна со своими настройками

        # Подключаем события перетаскивания
        self.pickFileButton.dragEnterEvent = self.dragEnterEvent
        self.pickFileButton.dropEvent = self.dropEvent
        # Добавляем обработку событий перетаскивания файлов
        self.filePicked.setAcceptDrops(True)

        # ДОБАВЛЕНИЕ ЭЛЕМЕНТОВ В ГРИД #
        grid = QGridLayout()
        grid.setSpacing(0)
        self.setLayout(grid)
        grid.setContentsMargins(62, 0, 0, 0)
        grid.setColumnStretch(0, 1)  # Установить вес (stretch) для первого столбца
        grid.setColumnStretch(1, 1)  # Установить вес для второго столбца
        grid.setColumnStretch(2, 2)  # Установить вес для третьего столбца
        self.setStyleSheet("background-color: #0074BA;")

        grid.addWidget(self.title, 0, 0)  # QLabel("Проверка файла по ГОСТам')
        grid.addWidget(self.mainWindow_button, 0, 2)  # QPushButton("Проверить по своим параметрам")
        grid.addWidget(self.labelGost, 1, 0)  # QLabel("Выберите ГОСТ")
        grid.addWidget(self.gostPicked, 5, 0)  # QLabel(self) выбранный гост
        grid.addWidget(self.elementPicked, 6, 0)  # QLabel(self) выбранный элемент
        grid.addWidget(self.choiceGost, 2, 0)  # QComboBox(self)
        grid.addWidget(self.choiceElement, 3, 0)  # QComboBox(self)
        grid.addWidget(self.fontStyleLabel, 2, 1)  # QLabel("Шрифт:")
        grid.addWidget(self.fontSizeLabel, 3, 1)  # QLabel("Размер шрифта:")
        grid.addWidget(self.paragraphIndentLabel, 4, 1)  # QLabel("Абзацный отступ:")
        grid.addWidget(self.intervalLabel, 5, 1)  # QLabel("Межстрочный интервал:")
        grid.addWidget(self.alignmentLabel, 6, 1)  # QLabel("Выравнивание:")
        grid.addWidget(self.styleLabel, 7, 1)  # QLabel("Название стиля:")
        grid.addWidget(self.pickFileButton, 8, 0, 1, 2)  # QPushButton("Выбрать файл")
        grid.addWidget(self.filePicked, 9, 0, 1, 2)  # QLabel("") выбранный файл
        grid.addWidget(self.checkFile, 10, 0, 1, 2)  # QPushButton("Проверить файл")
        grid.addWidget(self.answer, 1, 2, 10, 1)  # QScrollArea(self)
        grid.addWidget(self.plain_text, 1, 2, 9, 1)  # QPlainTextEdit()
        grid.addWidget(self.downloadFile, 10, 2, 1, 1)  # QPushButton("Скачать проверенный файл")




    def get_params_from_ghost(self, index):
        gost = self.choiceGost.itemText(index)
        self.selected_gost = gost if gost else self.selected_gost
        self.gostPicked.setText(self.selected_gost)
        print("selected_gost: ", self.selected_gost)

    def get_params_from_element(self, index):
        self.object = FileManager(1, docx.Document('../test2.docx'), 'tur', gost=self.selected_gost, doc_rej=False)
        asyncio.run(self.object.update_params_from_gost())
        self.selected_element = self.choiceElement.itemText(index)
        self.elementPicked.setText(self.selected_element)
        styles_elem = self.object.styles
        print("styles: ", styles_elem)

        self.selected_element = self.choiceElement.itemText(index)
        self.elementPicked.setText(self.selected_element)
        print("selected_element: ", self.selected_element)

        self.style_name = styles_elem[self.selected_element].style_name
        self.alignment = styles_elem[self.selected_element].alignment
        self.indent = styles_elem[self.selected_element].indent
        self.interval = styles_elem[self.selected_element].interval
        self.fname = styles_elem[self.selected_element].fname
        self.fsize = styles_elem[self.selected_element].fsize

        self.fontStyleLabel.setText(f"Шрифт: {self.fname}")
        self.fontSizeLabel.setText(f"Размер шрифта: {self.fsize}")
        self.paragraphIndentLabel.setText(f"Абзацный отступ: {self.indent}")
        self.intervalLabel.setText(f"Межстрочный интервал: {self.interval}")
        self.alignmentLabel.setText(f"Выравнивание: {self.alignment}")
        self.styleLabel.setText(f"Название стиля: {self.style_name}")


    # def loadGostValues(self):
    #     with open('GOSTs.json', 'r', encoding='utf-8') as file:
    #         gost_values = json.load(file)
    #     return gost_values

    def open_main_window(self):
        from mainWindow import MainWindow
        self.main_window = MainWindow()
        self.main_window.show()
        self.close()

    def pickFileButton_Clicked(self):
        filename, filetype = QFileDialog.getOpenFileName(self,
                                                         "Выбрать файл",
                                                         '.',
                                                         'Word files (*.docx)')
        if filename == '':
            self.filePicked.setText('Файл не выбран.')
            self.pathFile = ''
        else:
            self.pathFile = filename
            filename = filename.split('/')[-1]
            self.filePicked.setText(filename)

    def checkFile_Clicked(self):
        print(self.pathFile)
        if self.pathFile == '':
            try:
                self.notSelectFile = QMessageBox()
                self.notSelectFile.setStandardButtons(QMessageBox.Ok | QMessageBox.Cancel)
                self.notSelectFile.setText('Вы не выбрали файл!')
                self.notSelectFile.setWindowTitle('Ошибка!')
                self.notSelectFile.setIcon(QMessageBox.Warning)
                res = self.notSelectFile.exec()
            except Exception as e:
                pass
        else:
            self.selected_gost = self.gostPicked.text()
            print(self.selected_gost)
            # print(type(self.selected_gost))
            self.fileName = self.pathFile.split('/')[-1]
            self.path = f'../{self.fileName}'
            print(self.path)
            obj = FileManager(1, docx.Document(self.path), 'tur', gost=self.selected_gost, doc_rej=False)
            errors = asyncio.run(obj.is_correct_document())
            self.plain_text.clear()
            self.plain_text.setPlainText(errors)



    def save_ready_file(self):
        obj2 = FileManager(1, docx.Document(self.path), 'tur', gost=self.selected_gost, doc_rej=True)
        asyncio.run(obj2.is_correct_document())
        options = QFileDialog.Options()
        default_name = "_ready_file.docx"
        file_name, _ = QFileDialog.getSaveFileName(None, "Сохранить файл", default_name, "Word files (*.docx)",
                                                   options=options)
        if file_name:
            obj2.user_file.save(file_name)


    def dragEnterEvent(self, event):
        mime_data = event.mimeData()
        if mime_data.hasUrls() and mime_data.urls()[0].isLocalFile():
            event.acceptProposedAction()

    def dropEvent(self, event):
        mime_data = event.mimeData()
        if mime_data.hasUrls() and mime_data.urls()[0].isLocalFile():
            file_path = mime_data.urls()[0].toLocalFile()
            filename = file_path.split('/')[-1]
            # filename = file_path
            self.filePicked.setText(filename)
            self.pathFile = file_path
            event.acceptProposedAction()



if __name__ == '__main__':
    app = QApplication(sys.argv)
    window = SecondWindow(None)
    window.show()
    sys.exit(app.exec_())


# if __name__ == '__main__':
#     app = QApplication([])
#     window = SecondWindow(None)
#     # window = SecondWindow(None)
#     window.show()
#     app.exec_()
