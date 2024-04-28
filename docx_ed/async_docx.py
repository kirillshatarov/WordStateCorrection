import asyncio
import os
import style_parser as sp
import docx
import docx_ed.cfg as c

from typing import Any
from dataclasses import dataclass
from docx_ed.file_reader import FileReader
from docx_ed.gen_template import Template


def painter(paragraph: docx, errors: list[tuple]):
    for color, comment in errors:
        paint(paragraph, color, comment)


def paint(paragraph: docx, color: str, comment: str):
    paragraph.add_comment(comment)
    for word in paragraph.runs:
        word.font.color.rgb = c.colors[color]


def join_numbers(numbers):
    if not numbers:
        return ""
    numbers = list(map(int, numbers))
    ranges = []
    start = end = numbers[0]

    for num in numbers[1:]:
        if num == end + 1:
            end = num
        else:
            if start == end:
                ranges.append(str(start))
            else:
                ranges.append(f"{start}-{end}")
            start = end = num

    if start == end:
        ranges.append(str(start))
    else:
        ranges.append(f"{start}-{end}")

    return ", ".join(ranges)


@dataclass
class StyleStorage:
    style_name: str
    alignment: docx.enum.text.WD_ALIGN_PARAGRAPH
    indent: Any
    interval: Any
    fsize: Any
    fname: str

    def __str__(self):
        return self.style_name


class FileManager:
    def __init__(self, user_id: int, docx_: docx.Document, name: str, doc_rej: bool = False,
                 gost=None):
        self.gost = gost
        self.user_id = user_id
        self.user_file = docx_
        self.user_file_name = name
        self.alignment = None
        self.indent = None
        self.interval = None
        self.styles = {}
        self.last_style = None
        self.doc_rej = doc_rej

    @staticmethod
    def msg_errors(errors) -> str:
        errors = list(errors)
        if errors:
            comments = []
            main_err = ''
            for main_err, i in errors:
                comments.append(i)
            ans = main_err[1][1]
            abzac = join_numbers(comments)
            ans += f'\nВ абзацах: \n{abzac}\n\n\t'
            return ans
        return "Все соответствует госту"

    def lineal_is_choosen(self, function_name) -> str:
        paragraphs = self.user_file.paragraphs
        errors = []
        for i, paragraph in enumerate(paragraphs):
            errors.append((self.get_f_dict()[function_name](paragraph), i))
        return self.msg_errors(filter(lambda tup: tup[0][0] is True, errors))

    async def update_params_from_gost(self):
        user_gosts = FileReader.get_user_gosts().keys()
        pre_gosts = FileReader.get_actual_pre_gosts().keys()
        if self.gost in user_gosts or self.gost in pre_gosts:
            fl = FileReader(self.gost + '.json')
            is_user_gost = False
            if self.gost in pre_gosts:
                gost_dicts = await fl.read_file_from_pre()

            else:
                gost_dicts = await fl.read_file_from_user()
                is_user_gost = True
            alignment_settings = c.setter_gost
            name = None
            for gd_name in gost_dicts:
                if name is None:
                    name = gost_dicts[gd_name]
                    continue
                style_states = gost_dicts[gd_name]
                if len(style_states) == 0: continue
                alignment = c.templ_sel_gost[style_states['alignment']] if is_user_gost else style_states['alignment']
                style = StyleStorage(
                    gd_name,
                    alignment_settings[alignment],
                    style_states['indent'],
                    style_states['interval'],
                    style_states['font-size'],
                    style_states['font-style']
                )

                self.styles[gd_name] = style
            return True
        return False

    def get_f_dict(self):
        return {
            'font-size': self.is_correct_font_size,
            'font-style': self.is_correct_font_style,
            'alignment': self.is_correct_alignment,
            'indent': self.is_correct_indent,
            'interval': self.is_correct_interval
        }

    def is_correct_font_size(self, paragraph) -> tuple:
        error = (False, ('pink', 'All-okey'))
        if self.last_style is not None:
            fsize = self.styles[self.last_style].fsize
        else:
            fsize = self.fsize
        if fsize is None: return error
        for run in paragraph.runs:
            font_size = run.font.size.pt if run.font.size else None
            if font_size is None: return error
            if isinstance(fsize, list):
                left, right = map(float, fsize)
                if not (left <= font_size <= right):
                    error = (True, ('pink', c.exceptions['font-size'] + '-'.join(fsize)))
                    break
            else:
                if font_size != float(fsize):
                    error = (True, ('pink', c.exceptions['font-size'] + str(fsize)))
                    break
        return error

    def is_correct_font_style(self, paragraph):
        error = (False, ('red', 'All-okey'))

        if self.last_style is not None:
            fname = self.styles[self.last_style].fname
        else:
            fname = self.fname

        if fname is None: return error
        for run in paragraph.runs:
            font_style = run.font.name
            if font_style:
                if isinstance(fname, list):
                    if font_style not in fname:
                        error = (
                            True, ('red', c.exceptions['font-style'] + '-'.join(fname)))
                else:
                    if font_style not in fname:
                        error = (True, ('red', c.exceptions['font-style'] + str(fname)))
        return error

    def is_correct_alignment(self, paragraph):
        error = (False, ('green', 'All-okey'))
        doc_alignment = paragraph.alignment
        doc_alignment = doc_alignment if doc_alignment else paragraph.style.paragraph_format.alignment

        if self.last_style is not None:
            alignment = self.styles[self.last_style].alignment
        else:
            alignment = self.alignment

        if alignment is None: return error

        if doc_alignment:
            if paragraph.text.strip():
                if alignment != doc_alignment:
                    error = (True, ('green', c.exceptions['alignment'] + str(alignment)))
        return error

    def is_correct_interval(self, paragraph):
        doc_interval = paragraph.style.paragraph_format.line_spacing
        error = (False, ('yellow', 'All-okey'))
        if doc_interval is None: return error
        doc_interval = round(doc_interval, 2)

        if self.last_style is not None:
            interval = self.styles[self.last_style].interval
        else:
            interval = self.interval

        if interval is None: return error
        if isinstance(interval, list):
            left, right = map(float, interval)
            if not (left <= doc_interval <= right):
                error = (True, ('yellow', c.exceptions['line_spacing'] + '-'.join(interval)))

        else:
            if doc_interval != float(interval):
                error = (True, ('yellow', c.exceptions['line_spacing'] + str(interval)))
        return error

    def is_correct_indent(self, paragraph):
        error = (False, ('blue', 'All-okey'))
        if self.last_style is not None:
            indent = self.styles[self.last_style].indent
        else:
            indent = self.indent

        if indent is None: return error

        if paragraph.paragraph_format.first_line_indent is not None:
            doc_indent = round(paragraph.paragraph_format.first_line_indent.cm, 2)
            if isinstance(indent, list):
                left, right = map(float, indent)
                if not (left <= doc_indent <= right):
                    error = (True, ('blue', c.exceptions['indent'] + '-'.join(indent)))
            else:
                if doc_indent != float(indent):
                    error = (True, ('blue', c.exceptions['indent'] + str(indent)))
        return error

    async def is_correct_document(self):
        if await self.update_params_from_gost():
            errors = {'alignment': [],
                      'line_spacing': [],
                      'indent': [],
                      'font-size': [],
                      'font-style': []
                      }

            first_page = False  # Скипает первую страницы до Введения у отчетов
            for i, paragraph in enumerate(self.user_file.paragraphs):
                self.last_style = 'main_text'
                if first_page:
                    if sp.is_Heading(paragraph.style.name):
                        first_page = False
                    else:
                        continue

                if sp.is_picture_or_figure(paragraph.style.name): continue
                if sp.is_Heading(paragraph.style.name): self.last_style = 'heading'
                if sp.is_listing(paragraph): continue

                actual_font_size_info = self.is_correct_font_size(paragraph)  # Размер шрифта
                actual_font_style_info = self.is_correct_font_style(paragraph)  # Стиль шрифта
                actual_alignment_info = self.is_correct_alignment(paragraph)  # Выравнивание
                actual_interval_info = self.is_correct_interval(paragraph)  # Межстрочный интервал
                actual_indent_info = self.is_correct_indent(paragraph)  # Отступы
                reports = [actual_font_size_info,
                           actual_font_style_info,
                           actual_alignment_info,
                           actual_interval_info,
                           actual_indent_info]
                err_filter = filter(lambda tup: tup[0] is True, reports)

                if err_filter:
                    err_comments = [err[1] for err in err_filter]
                    if self.doc_rej:
                        painter(paragraph, err_comments)
                    else:
                        for color, comment in err_comments:
                            errors[c.color_exceptions[color]].append(comment + f'\n На строке {i}')
            return self.answer(errors)
        return False

    def answer(self, errors: dict = None):
        if self.doc_rej:
            self.saver()
        else:
            if self.gost in FileReader.get_user_gosts():
                answer = f'Проверка госта из примера\n'
            else:
                answer = f'Проверка госта {self.gost}\n'
            for keyh in errors:
                reason, spec = '', ''
                if errors[keyh]:
                    abzac = []
                    for er in errors[keyh]:
                        reason, fix = er.split(':')
                        spec, strings = fix.split('\n')
                        abzac.append(strings.split()[-1])
                    abzac = join_numbers(abzac)
                    answer += f' Проблемы возникли с {keyh}: \n\t{reason.strip()} {spec}\n\t  В абзацах: \n{abzac}\n\n\t'
            if answer != f'Проверка госта {self.gost}\n':
                return answer
            else:
                return 'Все соотвествует ГОСТу'

    def saver(self):
        self.user_file.save(f'../files/edited_Docx/{self.user_id}_ready_file.docx')

    def close(self):
        os.remove(self.user_file_name)


if __name__ == '__main__':
    template = Template(1, docx.Document('../test.docx'))
    template.writeTemplates(template.generate_gost())
    obj = FileManager(1, docx.Document('../test2.docx'), 'tur', gost="new_gost", doc_rej=False)
    print(asyncio.run(obj.is_correct_document()))
