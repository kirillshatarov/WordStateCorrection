import os
import docx
import docx_ed.cfg as c
from docx_ed.file_reader import FileReader
from docx.shared import Pt

def painter(paragraph: docx, errors: list[tuple]):
    for color, comment in errors:
        paint(paragraph, color, comment)


def paint(paragraph: docx, color: str, comment: str):
    paragraph.add_comment(comment)
    for word in paragraph.runs:
        word.font.color.rgb = c.colors[color]


class FileManger:
    def __init__(self, user_id: int, docx_: docx.Document, name: str, doc_rej: bool = False,
                 gost=None):
        self.gost = gost
        self.user_id = user_id
        self.user_file = docx_
        self.user_file_name = name
        self.alignment = None
        self.indent = None
        self.interval = None
        self.fsize = None
        self.fname = None
        self.doc_rej = doc_rej

    @staticmethod
    def msg_errors(errors: list) -> str:
        if errors:
            return "\n".join(errors)
        return "Все соответствует госту"

    def is_fully_alignment(self) -> str:
        paragraphs = self.user_file.paragraphs
        errors = []
        for i, paragraph in enumerate(paragraphs):
            if paragraph.alignment != self.alignment:
                errors.append(f"Не соответствие госту в {i + 1} строке")
        return self.msg_errors(errors)

    def is_correct_font_size(self):
        paragraphs = self.user_file.paragraphs
        errors = []
        for i, paragraph in enumerate(paragraphs):
            font_size = paragraph.style.font.size
            if font_size:
                font_size = font_size.pt
                if isinstance(self.fsize,list):
                    left, right = map(int, self.fsize)
                    if left <= font_size <= right:
                        errors.append(f"Размер шрифта не соответствие госту в {i + 1} строке")
                else:
                    if font_size != self.fsize:
                        errors.append(f"Размер шрифта не соответствие госту в {i + 1} строке")
        return self.msg_errors(errors)

    def is_correct_font_style(self):
        paragraphs = self.user_file.paragraphs
        errors = []
        for i, paragraph in enumerate(paragraphs):
            font_style = paragraph.style.font.name
            if font_style:
                if isinstance(self.fname, list):
                    if font_style in self.fname:
                        errors.append(f"Стиль шрифта не соответствие госту в {i + 1} строке")
                else:
                    if font_style != self.fname:
                        errors.append(f"Стиль шрифта не соответствие госту в {i + 1} строке")
        return self.msg_errors(errors)

    def is_correct_indents(self):
        paragraphs = self.user_file.paragraphs
        errors = []
        for i, paragraph in enumerate(paragraphs):
            if paragraph.style.name not in c.TITLE.values():
                if len(paragraph._element.xpath('./w:pPr/w:numPr')) > 0:
                    continue
                if paragraph.paragraph_format.first_line_indent is not None:
                    doc_indent = round(paragraph.paragraph_format.first_line_indent.cm, 2)
                    if isinstance(self.indent, list):
                        left, right = self.indent
                        if not (left <= doc_indent <= right):
                            errors.append(f'Абзац: "{paragraph.text[:25]}" оформлен неверно. Его отступ '
                                          f'составляет: {doc_indent} см')
                    else:
                        if doc_indent != float(self.indent):
                            errors.append(f'Абзац: "{paragraph.text[:25]}" оформлен неверно. Его отступ '
                                          f'составляет: {doc_indent} см')
        return self.msg_errors(errors)

    def is_correct_line_spacing(self):
        lines = [paragraph.paragraph_format.line_spacing for paragraph in self.user_file.paragraphs if
                 paragraph.style.name not in c.TITLE.values()]
        errors = []
        for i, p_ in enumerate(lines):
            if isinstance(self.interval, list):
                if self.interval[0] <= p_ <= self.interval[1]:
                    errors.append(f"Не соответствие госту в {i + 1} строке")
            else:
                if p_ != self.interval:
                    errors.append(f"Не соответствие госту в {i + 1} строке")
        return self.msg_errors(errors)

    def get_params_from_ghost(self):
        if self.gost in FileReader.get_files().keys():
            params = FileReader(self.gost + '.json').read_file()
            self.alignment = c.setter_gost[params['alignment']]
            self.indent = params['paragraph-indent']
            self.interval = params['interval']
            self.fname = params['font-style']
            self.fsize = params['font-size']
            return True
        return False

    def is_correct_document(self):
        if self.get_params_from_ghost():
            errors = {'alignment': [],
                      'line_spacing': [],
                      'indent': [],
                      'font-size':[],
                      'font-style':[]
                      }
            for i, paragraph in enumerate(self.user_file.paragraphs):
                if not (paragraph.style.name.startswith('Heading') and
                        paragraph.style.name.startswith('Subheading')):
                    # Для списков
                    par_errors = []
                    if len(paragraph._element.xpath('./w:pPr/w:numPr')) > 0:
                        continue
                    #Размер шрифта
                    font_size = paragraph.style.font.size
                    if font_size:
                        font_size = font_size.pt
                        if isinstance(self.fsize, list):
                            left,right = map(float,self.fsize)
                            if left <= font_size <= right:
                                comment = c.exceptions['font-size'] + '-'.join(self.fsize)
                                par_errors.append(('pink', comment))
                        else:
                            if font_size != float(self.fsize):
                                comment = c.exceptions['font-size'] + str(self.fsize)
                                par_errors.append(('pink', comment))
                    #Стиль шрифта
                    font_style = paragraph.style.font.name
                    if font_style:
                        if isinstance(self.fname,list):
                            if font_style not in self.fname:
                                comment = c.exceptions['font-style'] + '-'.join(self.fname)
                                par_errors.append(('red', comment))
                        else:
                            if font_style not in self.fname:
                                comment = c.exceptions['font-style'] + str(self.fname)
                                par_errors.append(('red', comment))

                    # Выравнивание
                    alignment = paragraph.alignment
                    if alignment:
                        if alignment != self.alignment:
                            comment = c.exceptions['alignment'] + str(self.alignment)
                            par_errors.append(('green', comment))

                    # Межстрочный интервал
                    interval = paragraph.paragraph_format.line_spacing
                    if interval:
                        if isinstance(self.interval, list):
                            left, right = map(float, self.indent)
                            if left <= interval <= right:
                                comment = c.exceptions['line_spacing'] + '-'.join(self.interval)
                                par_errors.append(('yellow', comment))

                        else:
                            if interval != self.interval:
                                comment = c.exceptions['line_spacing'] + str(self.interval)
                                par_errors.append(('yellow', comment))

                    # абзацы скипаем
                    if paragraph.paragraph_format.first_line_indent is not None:
                        # Отступ
                        doc_indent = round(paragraph.paragraph_format.first_line_indent.cm, 2)
                        if isinstance(self.indent, list):
                            left, right = map(float, self.indent)
                            if not (left <= doc_indent <= right):
                                comment = c.exceptions['indent'] + '-'.join(self.indent)
                                par_errors.append(('blue', comment))
                        else:
                            if doc_indent != float(self.indent):
                                comment = c.exceptions['indent'] + str(self.indent)
                                par_errors.append(('blue', comment))

                    if par_errors:
                        if self.doc_rej:
                            painter(paragraph, par_errors)
                        else:
                            for color, comment in par_errors:
                                errors[c.color_exceptions[color]].append(comment + f'\n На строке {i}')
            return self.answer(errors)
        return False

    def answer(self, errors: dict = None):
        if self.doc_rej:
            self.saver()
        else:
            answer = ''
            for keyh in errors:
                if errors[keyh]:
                    answer += f'Проблемы возникли с {keyh}: \n' + '\n'.join(errors[keyh]) + '\n'
            if answer:
                return answer
            else:
                return 'Все соотвествует ГОСТу'

    def saver(self):
        self.user_file.save(f'../files/edited_Docx/{self.user_id}_ready_file.docx')

    def close(self):
        os.remove(self.user_file_name)


if __name__ == '__main__':
    obj = FileManger(1, docx.Document('../test2.docx'), 'tur', gost="2.105-2019", doc_rej=True)
    obj.is_correct_document()
