import logging
import os

import aiofiles
from aiogram import Bot, Dispatcher, types
from aiogram.dispatcher import FSMContext
from aiogram.dispatcher.filters.state import StatesGroup, State
from aiogram.utils import executor
from docx import Document
from docx_ed import async_docx as dc, file_reader as fl
from aiogram.contrib.fsm_storage.mongo import MongoStorage

from bot import cfg as c

# Initialize the bot and dispatcher
bot = Bot(token=c.TOKEN)
logging.basicConfig(level=logging.INFO)
# Establish connection to MongoDB
"""mongo_client = MongoClient('localhost', 27017)
mongo_db = mongo_client['your_database_name']
storage = MongoStorage(mongo_db, 'states_data')"""
storage = MongoStorage(uri="mongodb://localhost:27017")
dp = Dispatcher(bot, storage=storage)


@dp.message_handler(content_types=types.ContentTypes.TEXT, state='*', commands='reset')
async def resetData(message: types.Message, state: FSMContext = None):
    await message.answer('Все данные были удалены')
    stat = await state.reset_state()
    data = await state.reset_data()


class Form(StatesGroup):
    file_ = State()
    prechoose = State()
    gost = State()
    final_gost = State()
    choose1 = State()
    choose2 = State()
    start = State()


# Reply keyboard markup
m_i_key = types.ReplyKeyboardMarkup(resize_keyboard=True, one_time_keyboard=True).add("1.0", "1.25", "1.5", "2")
alig_key = types.ReplyKeyboardMarkup(resize_keyboard=True, one_time_keyboard=True).add(
    "по ширине", "по левому краю", "по правому краю", "по центру", "по умолчанию")
start_keys = types.ReplyKeyboardMarkup(resize_keyboard=True, one_time_keyboard=True, row_width=1).add(
    "Проверить конкретный гост", "Проверить отдельно", "Прекратить взаимодействие")
other_keys = types.ReplyKeyboardMarkup(resize_keyboard=True, one_time_keyboard=True).add(
    "Проверка выравнивания", "Проверка межстрочного интервала", "Проверка абзацных отступов",
    "Прекратить взаимодействие")
a_i_key = types.ReplyKeyboardMarkup(resize_keyboard=True, one_time_keyboard=True).add("1.0", "1.25", "1.5", "2", "3")
file_text_keys = types.ReplyKeyboardMarkup(resize_keyboard=True, one_time_keyboard=True, row_width=1).add(
    "Отправить файл", "Отправить текст", "Прекратить взаимодействие")


# Function to generate gost keys
def gost_keys():
    gkey = types.ReplyKeyboardMarkup(resize_keyboard=True, one_time_keyboard=True)
    for ke in fl.FileReader.get_files():
        gkey.add(ke)
    return gkey


gkeys = gost_keys()


@dp.message_handler(commands=["start"])
async def aboba(message: types.Message):
    await message.answer("Загружайте файл для проверки на соответствие госту")
    await Form.file_.set()


@dp.message_handler(state=Form.file_, content_types=types.ContentType.DOCUMENT)
async def handle_docs(message: types.Message, state: FSMContext):
    if message.document.mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
        doc_name = await message.document.download(destination_dir="../files/")
        async with state.proxy() as data:
            # Create a dictionary containing relevant data from FileManager
            data['doc_obj'] = {'user_id': message.chat.id,
                               'name': doc_name.name}

        await message.answer("Спасибо, ваш файл docx получен и обработан! Теперь отправьте, что вы хотите проверить",
                             reply_markup=start_keys)
        await Form.prechoose.set()
    else:
        await message.answer("Пожалуйста, отправьте файл в формате docx.")


@dp.message_handler(state=Form.prechoose)
async def process(message: types.Message, state: FSMContext):
    async with state.proxy() as data:
        if message.text == "Проверить конкретный гост":
            await message.answer("Спасибо, за выбор. Теперь выберите гост, который вы хотите проверить",
                                 reply_markup=gkeys)
            await Form.gost.set()
        elif message.text == "Проверить отдельно":
            await message.answer("Теперь отправьте, что вы хотите проверить", reply_markup=other_keys)
            await Form.choose1.set()
        elif message.text == "Прекратить взаимодействие":
            await state.finish()
        else:
            await message.answer(
                "Отправьте, что вы хотите проверить",
                reply_markup=start_keys)


@dp.message_handler(state=Form.gost)
async def process_gost(message: types.Message, state: FSMContext):
    async with state.proxy() as data:
        if message.text in fl.FileReader.get_files():
            data['doc_obj']['gost'] = message.text
            await message.answer("Выберите режим взаимодействия", reply_markup=file_text_keys)
            await Form.final_gost.set()
        else:
            await Form.start.set()
            await message.answer('Данного госта нет в базе', reply_markup=gkeys)


@dp.message_handler(state=Form.final_gost)
async def final_process_gost(message: types.Message, state: FSMContext):
    async with state.proxy() as data:
        documenter = dc.FileManager(user_id=message.chat.id,
                                    docx_=Document(data['doc_obj']['name']),
                                    name=data['doc_obj']['name'],
                                    gost=data['doc_obj']['gost'],
                                    doc_rej=False)

        if message.text == 'Отправить файл':
            documenter.doc_rej = True
            await documenter.is_correct_document()
            file_path = f"../files/edited_Docx/{message.chat.id}_ready_file.docx"
            async with aiofiles.open(file_path, 'rb') as file:
                await bot.send_document(chat_id=message.chat.id, document=file)
            os.remove(file_path)
            await message.answer("Спасибо за использование нашего бота!")
        elif message.text == 'Отправить текст':
            response = await documenter.is_correct_document()
            await message.answer(response)
        elif message.text == "Прекратить взаимодействие":
            await message.answer(
                "Спасибо за использование нашего бота! Для повторного использования вызовете вновь команду /start")
        else:
            await message.answer(
                "Отправьте, в каком виде вы хотите получить ответ",
                reply_markup=file_text_keys)
    documenter.close()
    await state.finish()


@dp.message_handler(state=Form.choose1)
async def process1(message: types.Message, state: FSMContext):
    async with state.proxy() as data:
        if message.text == "Проверка выравнивания":
            data['bot_rej'] = 0
            await message.answer("Спасибо за выбор. Теперь выберите требование", reply_markup=alig_key)
            await Form.choose2.set()
        elif message.text == "Проверка межстрочного интервала":
            data['bot_rej'] = 1
            await message.answer("Спасибо за выбор. Теперь выберите требование", reply_markup=m_i_key)
            await Form.choose2.set()
        elif message.text == "Проверка абзацных отступов":
            data['bot_rej'] = 2
            await message.answer("Спасибо за выбор. Теперь выберите требование", reply_markup=a_i_key)
            await Form.choose2.set()
        elif message.text == "Прекратить взаимодействие":
            documenter = dc.FileManager(user_id=message.chat.id,
                                        docx_=Document(data['doc_obj']['name']),
                                        name=data['doc_obj']['name'])
            documenter.close()
            await state.finish()
        else:
            await message.answer("Отправьте, что хотите проверить.", reply_markup=other_keys)


@dp.message_handler(state=Form.choose2)
async def process2(message: types.Message, state: FSMContext):
    async with state.proxy() as data:
        documenter = dc.FileManager(user_id=message.chat.id,
                                    docx_=Document(data['doc_obj']['name']),
                                    name=data['doc_obj']['name'])
        if data['bot_rej'] == 0:
            documenter.alignment = message.text
            await message.answer(documenter.lineal_is_correct_alignment(), reply_markup=None)
        elif data['bot_rej'] == 1:
            documenter.interval = float(message.text)
            await message.answer(documenter.lineal_is_correct_interval(), reply_markup=None)
        elif data['bot_rej'] == 2:
            documenter.indent = float(message.text)
            await message.answer(documenter.lineal_is_correct_indent(), reply_markup=None)

    await Form.choose1.set()
    await message.answer("Спасибо за использование нашего бота, вы можете выбрать другую функцию для вашего файла",
                         reply_markup=start_keys)


if __name__ == '__main__':
    executor.start_polling(dp, skip_updates=True)
